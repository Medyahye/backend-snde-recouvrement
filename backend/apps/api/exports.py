"""Génération des exports CSV et Word.

CSV : streaming via `csv` stdlib pour gérer des dizaines de milliers de lignes
sans charger en RAM.
Word : génération en mémoire via `python-docx`, retournée comme bytes.
"""
from __future__ import annotations

import csv
import io
from typing import Iterable

from django.http import StreamingHttpResponse
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, RGBColor

from apps.clients.models import Client
from apps.imports.models import FabImport
from apps.zones.models import Zone

from . import stats as stats_module


# --------------------------------------------------------------------------- #
# Helpers CSV (streaming)
# --------------------------------------------------------------------------- #


class _Echo:
    """Buffer pseudo-fichier pour le streaming CSV (cf. doc Django)."""

    def write(self, value):
        return value


def _format_mru(value) -> str:
    """Format MRU avec espaces millier : 1906305 → '1 906 305'."""
    if value is None:
        return ""
    return f"{int(value):,}".replace(",", " ")


def _format_score(value) -> str:
    """Format score : 0.3178 → '0,3178' (virgule décimale française)."""
    if value is None:
        return ""
    return f"{float(value):.4f}".replace(".", ",")


# --------------------------------------------------------------------------- #
# Export CSV — Zones
# --------------------------------------------------------------------------- #

CSV_ZONES_HEADERS = [
    "Rang",
    "Zone",
    "Centre",
    "Secteur",
    "Tournée",
    "Nb clients",
    "Nb entreprises",
    "Nb domestiques",
    "Score moyen",
    "Score max",
    "Anciennete moyenne (j)",
    "Solde total (MRU)",
    "Arriérés total (MRU)",
    "Priorité zone",
    "Priorité",
]


def _row_for_zone(z: Zone) -> list[str]:
    return [
        str(z.rang),
        z.zone_id,
        z.centre_nom,
        z.secteur,
        z.tournee,
        str(z.nb_clients),
        str(z.nb_entreprises),
        str(z.nb_domestiques),
        _format_score(z.score_moyen),
        _format_score(z.score_max),
        f"{z.anciennete_moyenne:.1f}".replace(".", ","),
        _format_mru(z.solde_total),
        _format_mru(z.arrieres_total),
        f"{z.priorite_zone:.2f}".replace(".", ","),
        z.priorite,
    ]


def stream_csv_zones(import_id: int) -> StreamingHttpResponse:
    writer = csv.writer(_Echo(), delimiter=";", quoting=csv.QUOTE_MINIMAL)
    zones = (
        Zone.objects.filter(import_ref_id=import_id).order_by("rang").iterator()
    )

    def rows():
        # BOM UTF-8 pour qu'Excel ouvre les accents proprement
        yield "﻿"
        yield writer.writerow(CSV_ZONES_HEADERS)
        for z in zones:
            yield writer.writerow(_row_for_zone(z))

    response = StreamingHttpResponse(rows(), content_type="text/csv; charset=utf-8")
    response["Content-Disposition"] = (
        f'attachment; filename="zones_import_{import_id}.csv"'
    )
    return response


# --------------------------------------------------------------------------- #
# Export CSV — Clients
# --------------------------------------------------------------------------- #

CSV_CLIENTS_HEADERS = [
    "Rang",
    "Référence",
    "Nom client",
    "Téléphone",
    "Adresse",
    "Activité",
    "Type",
    "Centre",
    "Secteur",
    "Tournée",
    "Zone",
    "Solde (MRU)",
    "Montant facture (MRU)",
    "Arriérés (MRU)",
    "Date facture",
    "Jours impayé",
    "Date dernier paiement",
    "Jours sans paiement",
    "Score",
    "Priorité",
]


def _row_for_client(c: Client) -> list[str]:
    return [
        str(c.rang),
        c.reference_abonnement,
        c.nom_client,
        c.telephone,
        c.adresse,
        c.activite_client,
        c.type_client,
        c.centre_nom,
        c.secteur_facturation,
        c.tournee_releve,
        c.zone,
        _format_mru(c.solde),
        _format_mru(c.montant_facture),
        _format_mru(c.arrieres),
        c.date_facture.isoformat() if c.date_facture else "",
        str(c.jours_impaye),
        c.date_dernier_paiement.isoformat() if c.date_dernier_paiement else "",
        str(c.jours_sans_paiement),
        _format_score(c.score_final),
        c.priorite,
    ]


def stream_csv_clients(import_id: int) -> StreamingHttpResponse:
    writer = csv.writer(_Echo(), delimiter=";", quoting=csv.QUOTE_MINIMAL)
    clients = (
        Client.objects.filter(import_ref_id=import_id).order_by("rang").iterator()
    )

    def rows():
        yield "﻿"
        yield writer.writerow(CSV_CLIENTS_HEADERS)
        for c in clients:
            yield writer.writerow(_row_for_client(c))

    response = StreamingHttpResponse(rows(), content_type="text/csv; charset=utf-8")
    response["Content-Disposition"] = (
        f'attachment; filename="clients_import_{import_id}.csv"'
    )
    return response


# --------------------------------------------------------------------------- #
# Export Word
# --------------------------------------------------------------------------- #


def _set_table_header(row, fill_hex="1F4E79"):
    """Style bleu marine SNDE pour l'entête d'un tableau."""
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
        # Fond bleu marine via XML brut (python-docx n'expose pas de helper)
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml

        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_kv(doc, label: str, value: str):
    """Ajoute un paragraphe 'Label : valeur' avec le label en gras."""
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label} : ")
    run_label.bold = True
    p.add_run(str(value))


def build_word_report(import_id: int) -> bytes:
    """Génère un rapport Word (docx) pour un import donné. Retourne les bytes."""
    imp = FabImport.objects.get(id=import_id)
    kpis = stats_module.compute_kpis(import_id)
    top_zones = list(
        Zone.objects.filter(import_ref_id=import_id).order_by("rang")[:20]
    )
    top_clients = list(
        Client.objects.filter(import_ref_id=import_id).order_by("rang")[:50]
    )

    doc = Document()

    # --- En-tête ---
    title = doc.add_heading("Rapport de Recouvrement SNDE", level=0)
    title.alignment = 1  # centré

    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run(
        f"Import du {imp.file_date.strftime('%d/%m/%Y')} — "
        f"généré le {imp.uploaded_at.strftime('%d/%m/%Y %H:%M')}"
    )
    run.font.size = Pt(10)
    run.italic = True

    doc.add_paragraph()

    # --- Section : Synthèse KPI ---
    doc.add_heading("1. Synthèse", level=1)
    _add_kv(doc, "Date du fichier", imp.file_date.strftime("%d/%m/%Y"))
    _add_kv(doc, "Statut", imp.get_status_display())
    _add_kv(doc, "Lignes brutes traitées", _format_mru(imp.nb_lines_total or 0))
    _add_kv(doc, "Clients retenus (relance=1)", _format_mru(kpis["totaux"]["nb_clients"]))
    _add_kv(doc, "Zones agrégées", str(kpis["totaux"]["nb_zones"]))
    _add_kv(doc, "Solde total à recouvrer", f"{_format_mru(kpis['totaux']['solde_total'])} MRU")
    _add_kv(doc, "Arriérés cumulés", f"{_format_mru(kpis['totaux']['arrieres_total'])} MRU")
    _add_kv(doc, "Score moyen global", _format_score(kpis["totaux"]["score_moyen"]))
    _add_kv(
        doc,
        "Répartition zones",
        (
            f"{kpis['zones_par_priorite']['Haute']} Haute / "
            f"{kpis['zones_par_priorite']['Moyenne']} Moyenne / "
            f"{kpis['zones_par_priorite']['Faible']} Faible "
            f"({kpis['zones_par_priorite']['pct_haute']}% en priorité Haute)"
        ),
    )
    if kpis["top_zone"]:
        tz = kpis["top_zone"]
        _add_kv(
            doc,
            "Zone n°1",
            f"{tz['zone_id']} — {tz['nb_clients']} clients, "
            f"score moyen {_format_score(tz['score_moyen'])}, "
            f"{_format_mru(tz['solde_total'])} MRU",
        )

    doc.add_paragraph()

    # --- Section : Top 20 zones ---
    doc.add_heading("2. Top 20 zones prioritaires", level=1)
    table = doc.add_table(rows=1, cols=8)
    table.style = "Light Grid Accent 1"
    headers = [
        "Rang",
        "Zone",
        "Centre",
        "Clients",
        "Score moy.",
        "Solde total (MRU)",
        "Arriérés (MRU)",
        "Priorité",
    ]
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    _set_table_header(hdr)

    for z in top_zones:
        row = table.add_row().cells
        row[0].text = str(z.rang)
        row[1].text = z.zone_id
        row[2].text = z.centre_nom
        row[3].text = str(z.nb_clients)
        row[4].text = _format_score(z.score_moyen)
        row[5].text = _format_mru(z.solde_total)
        row[6].text = _format_mru(z.arrieres_total)
        row[7].text = z.priorite

    doc.add_paragraph()

    # --- Section : Top 50 clients ---
    doc.add_heading("3. Top 50 clients prioritaires", level=1)
    table = doc.add_table(rows=1, cols=8)
    table.style = "Light Grid Accent 1"
    headers = [
        "Rang",
        "Référence",
        "Nom",
        "Téléphone",
        "Type",
        "Zone",
        "Solde (MRU)",
        "Score",
    ]
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    _set_table_header(hdr)

    for c in top_clients:
        row = table.add_row().cells
        row[0].text = str(c.rang)
        row[1].text = c.reference_abonnement
        row[2].text = c.nom_client
        row[3].text = c.telephone
        row[4].text = c.type_client
        row[5].text = c.zone
        row[6].text = _format_mru(c.solde)
        row[7].text = _format_score(c.score_final)

    # --- Pied : avertissement méthodologique ---
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "Note : le score combine 4 composantes (montant 40%, ancienneté 25%, "
        "historique 20%, arriérés 15%) avec un coefficient de 1.20 pour les "
        "clients Entreprise. La priorité zone = score moyen × nombre de clients. "
        "Les seuils Haute/Moyenne/Faible sont calculés par quantiles 75/50."
    )
    run.font.size = Pt(8)
    run.italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
