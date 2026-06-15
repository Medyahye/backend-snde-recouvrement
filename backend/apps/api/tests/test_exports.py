"""Tests des endpoints exports CSV + Word (étape 4.4)."""
import io

from django.urls import reverse


# --- CSV zones ---

def test_export_csv_zones_default(authed_client, fab_import):
    res = authed_client.get(
        reverse("export_csv"), {"import_id": fab_import.id}
    )
    assert res.status_code == 200
    assert "text/csv" in res["Content-Type"]
    assert "attachment" in res["Content-Disposition"]
    assert f"zones_import_{fab_import.id}" in res["Content-Disposition"]

    body = b"".join(res.streaming_content).decode("utf-8")
    # BOM + header + 6 lignes de zones
    assert "Rang;Zone;Centre" in body
    assert "KIFFA2_18_11" in body
    # Format MRU avec espaces
    assert "100 000" in body or "100000" not in body  # soit espacé, soit absent


def test_export_csv_zones_explicit_type(authed_client, fab_import):
    res = authed_client.get(
        reverse("export_csv"), {"import_id": fab_import.id, "type": "zones"}
    )
    assert res.status_code == 200


# --- CSV clients ---

def test_export_csv_clients(authed_client, fab_import):
    res = authed_client.get(
        reverse("export_csv"), {"import_id": fab_import.id, "type": "clients"}
    )
    assert res.status_code == 200
    body = b"".join(res.streaming_content).decode("utf-8")
    assert "Référence" in body
    assert "MAHFOUDH AMI" in body
    assert "R001" in body


# --- Erreurs CSV ---

def test_export_csv_invalid_type(authed_client, fab_import):
    res = authed_client.get(
        reverse("export_csv"),
        {"import_id": fab_import.id, "type": "invalid"},
    )
    assert res.status_code == 400


def test_export_csv_unknown_import(authed_client, fab_import):
    res = authed_client.get(reverse("export_csv"), {"import_id": 999_999})
    assert res.status_code == 400


def test_export_csv_unauthenticated(db, fab_import):
    from rest_framework.test import APIClient
    api = APIClient()
    res = api.get(reverse("export_csv"), {"import_id": fab_import.id})
    assert res.status_code == 401


# --- Word ---

def test_export_word_returns_docx(authed_client, fab_import):
    res = authed_client.get(reverse("export_word"), {"import_id": fab_import.id})
    assert res.status_code == 200
    assert "wordprocessingml" in res["Content-Type"]
    # Le filename inclut la date au format AAAAMMJJ
    assert "Rapport_SNDE_20260505.docx" in res["Content-Disposition"]
    # Vérifie que le contenu est un vrai docx (commence par PK = signature ZIP)
    assert res.content.startswith(b"PK")


def test_export_word_content_is_readable(authed_client, fab_import):
    """Charge le docx généré pour vérifier qu'il est valide et contient les sections clés."""
    from docx import Document

    res = authed_client.get(reverse("export_word"), {"import_id": fab_import.id})
    assert res.status_code == 200

    doc = Document(io.BytesIO(res.content))
    # Concatène tout le texte pour chercher les sections
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Rapport de Recouvrement SNDE" in full_text
    assert "Synthèse" in full_text or "Synthese" in full_text
    assert "Top 20 zones" in full_text
    assert "Top 50 clients" in full_text

    # Au moins 2 tables (top zones + top clients)
    assert len(doc.tables) >= 2


def test_export_word_unknown_import(authed_client, fab_import):
    res = authed_client.get(reverse("export_word"), {"import_id": 999_999})
    assert res.status_code == 400
